from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\Admin\Documents\roboparty_train_ljc")
OUT = ROOT / "研究" / "首批10篇论文方法综述与课题应用建议.docx"

BLUE = "2E5E8C"
DARK = "17324D"
MUTED = "5B6573"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
GRID = "B8C4D1"
WHITE = "FFFFFF"
RED = "9B1C1C"
GOLD = "7A5A00"


def set_run_font(run, east="宋体", latin="Calibri", size=10.5, bold=None, color=None, italic=None):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths, indent=120):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[i] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_para(doc, text="", style=None, bold_prefix=None, color=None, align=None, after=6, first_indent=True):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if first_indent and style is None:
        p.paragraph_format.first_line_indent = Pt(21)
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True, color=color)
        set_run_font(p.add_run(text[len(bold_prefix):]), color=color)
    else:
        set_run_font(p.add_run(text), color=color)
    return p


def add_label_para(doc, label, text, color=DARK, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    set_run_font(p.add_run(label), bold=True, color=color)
    set_run_font(p.add_run(text))
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.38 if level == 0 else 0.65)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_run_font(p.add_run("• "))
    set_run_font(p.add_run(text))
    return p


def create_numbering(doc):
    """Create a fresh decimal list so independent sections restart at 1."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(p, num_id):
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(num)
    p_pr.append(num_pr)


def add_number(doc, text, number):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_run_font(p.add_run(f"{number}. "))
    set_run_font(p.add_run(text))
    return p


def add_callout(doc, label, text, fill=PALE, color=DARK):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9300], indent=120)
    set_table_borders(table, color=fill, size=4)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    set_run_font(p.add_run(label), bold=True, color=color)
    set_run_font(p.add_run(text), color=color)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_source_line(doc, authors, year, arxiv, local_name, venue=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    set_run_font(p.add_run("来源："), bold=True, color=MUTED, size=9.5)
    venue_text = f"；{venue}" if venue else ""
    set_run_font(p.add_run(f"{authors}（{year}）{venue_text}；"), color=MUTED, size=9.5)
    add_hyperlink(p, f"arXiv:{arxiv}", f"https://arxiv.org/abs/{arxiv}")
    set_run_font(p.add_run(f"；本地文件：{local_name}"), color=MUTED, size=9.5)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.25)
section.bottom_margin = Cm(2.25)
section.left_margin = Cm(2.3)
section.right_margin = Cm(2.3)
section.header_distance = Cm(1.25)
section.footer_distance = Cm(1.25)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ("Title", 25, DARK, 0, 8),
    ("Subtitle", 13, MUTED, 0, 12),
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 11.5, DARK, 10, 5),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = name != "Subtitle"
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for list_style in ["List Bullet", "List Bullet 2", "List Number"]:
    st = styles[list_style]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.25

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hp.paragraph_format.space_after = Pt(0)
set_run_font(hp.add_run("课题文献方法综述 · 首批10篇"), east="微软雅黑", size=8.5, color=MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(0)
set_run_font(fp.add_run("第 "), size=9, color=MUTED)
add_field(fp, "PAGE")
set_run_font(fp.add_run(" 页"), size=9, color=MUTED)

# Editorial cover
for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
set_run_font(p.add_run("研究方法综述"), east="微软雅黑", size=12, bold=True, color=GOLD)

p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("基于不确定性感知的人形机器人\n鲁棒自适应运动控制"), east="微软雅黑", size=25, bold=True, color=DARK)

p = doc.add_paragraph(style="Subtitle")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("首批10篇相关论文的方法介绍、研究判断与 Atom-01 应用方案"), east="微软雅黑", size=13, color=MUTED)

for _ in range(3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("课题：RoboParty Atom-01 · Isaac Lab · rsl_rl/PPO · ROS2"), east="微软雅黑", size=10.5, bold=True, color=BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("依据《林金诚开题报告》整理 · 2026年8月17日"), size=10, color=MUTED)

doc.add_page_break()

doc.add_heading("1. 课题理解与本轮检索结论", level=1)
add_para(doc, "本课题拟面向复杂未知非结构化地形，研究人形机器人在视觉遮挡、点云飞点、时延、里程计漂移、感知盲区等条件下的鲁棒自适应运动控制。计划以 RoboParty Atom-01 为对象，在 Isaac Lab 中训练基于 rsl_rl/PPO 的单体控制策略，以局部高程图与本体状态作为多模态输入，经本体引导的交叉注意力进行连续软融合；同时用不确定性驱动的训练机制，使机器人在感知退化时降低速度、提高足端净空并增强姿态稳定，最终通过 MuJoCo Sim-to-Sim、ROS2、ONNX/TensorRT 和 elevation_mapping_cupy 完成实机验证。")
add_callout(doc, "本轮结论：", "最值得直接借鉴的不是某一篇论文的完整系统，而是四条可组合的方法线：Miki 等的带门控信念编码器、AME-2 的逐栅格不确定性地图、PIM/DreamWaQ 的内部模型与时序创新、RMA 的本体历史快速适应。它们可以合并成一个运行时单策略，而教师、特权 critic 和辅助预测头只在训练期使用。")

doc.add_heading("1.1 首批10篇论文与课题模块对应关系", level=2)
rows = [
    ("1", "RMA", "本体历史→隐式环境表征与快速适应", "核心"),
    ("2", "Robust Perceptive Locomotion", "门控信念状态、噪声课程、软融合", "核心"),
    ("3", "LocoTransformer", "视觉/本体 token 的跨模态注意力", "核心"),
    ("4", "Vision-Based Bipedal Locomotion", "双足高程图预测、视觉调制、延迟随机化", "核心"),
    ("5", "DreamWaQ", "本体时序内部模型、β-VAE、环境隐变量", "重要"),
    ("6", "Perceptive Internal Model", "人形机器人高程图内部模型与跨平台验证", "核心"),
    ("7", "No More Blind Spots", "全向双足视觉、DAgger、差分动作", "重要"),
    ("8", "AME-2", "不确定性地图、全局/局部注意力、联合RL蒸馏", "核心"),
    ("9", "Hybrid Autoencoder", "多传感器时序高程图重建", "补充"),
    ("10", "MARCH", "安全参考、CLF奖励、视觉学生策略", "重要"),
]
table = doc.add_table(rows=1, cols=4)
headers = ["序号", "论文简称", "可迁移的方法模块", "优先级"]
for i, h in enumerate(headers):
    set_cell_shading(table.rows[0].cells[i], LIGHT)
    p = table.rows[0].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(h), east="微软雅黑", size=9.5, bold=True, color=DARK)
repeat_header(table.rows[0])
for row in rows:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        p = cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(p.add_run(value), size=9.3)
set_table_geometry(table, [620, 2340, 5400, 940], indent=120)
set_table_borders(table)
doc.add_paragraph()

doc.add_heading("1.2 三项必须先澄清的方法学问题", level=2)
add_number(doc, "注意力权重或注意力熵不能自动等同于经过校准的感知不确定性。它可以作为特征或诊断量，但核心风险量应主要来自高程预测方差、传感器缺失/时延、时序创新残差，以及必要时的模型集成或 MC Dropout 分歧。", 1)
add_number(doc, "奖励函数只在训练阶段存在。若希望实机推理时根据感知质量主动变慢、抬高脚或增加稳定裕度，必须把风险量显式作为策略输入，或用它连续调节视觉残差分支/速度命令；不能只在训练代码中动态改奖励权重。", 2)
add_number(doc, "“单网络”应定义为运行时只有一个 actor，而不是禁止使用训练期教师、特权 critic、地图监督或辅助预测头。这样既保留论文创新中的单策略软退化，也能采用成熟的特权学习提高样本效率和实机成功率。", 3)

doc.add_page_break()
doc.add_heading("2. 十篇论文的方法介绍与课题应用", level=1)

# Paper 1
doc.add_heading("2.1 RMA：Rapid Motor Adaptation for Legged Robots", level=2)
add_source_line(doc, "A. Kumar, Z. Fu, D. Pathak, J. Malik", "2021", "2107.04034", "01_RMA_2107.04034.pdf", "Robotics: Science and Systems")
add_label_para(doc, "方法介绍。", "RMA 把运动策略拆成快速基策略和较慢的适应模块。训练时，基策略接收当前本体状态、上一时刻动作和 8 维环境隐变量 z；z 由可访问质量、摩擦、载荷、电机强度和局部地形等特权参数的环境编码器产生。部署时这些特权参数不可见，因此适应模块使用最近 0.5 s 的状态-动作历史，经 MLP 与 1D-CNN 直接回归 z，而不是显式辨识每个物理参数。为减小分布偏移，适应模块使用自身预测参与 rollout，再以仿真真值 z 监督，属于 DAgger 式在策略数据聚合。基策略约 100 Hz、适应模块约 10 Hz 异步运行。")
add_label_para(doc, "我的想法。", "RMA 最适合作为课题的“纯本体兜底”来源，而不是完整控制架构。它说明无需恢复可解释的摩擦或载荷参数，只要学习对动作有用的低维环境隐变量即可。对本课题而言，这个隐变量还能表征接触异常、地形软硬和执行器误差，并与视觉风险量形成互补。")
add_label_para(doc, "如何应用。", "在 Atom-01 的本体编码器后增加一个 0.4-0.6 s 历史窗口 TCN/1D-CNN，输出 z_prop；仿真中用质量、摩擦、电机强度、延迟、接触地形等特权量生成 z_priv，并以 MSE/对比损失监督 z_prop。让 z_prop 同时参与本体 Query、风险估计和纯本体动作主干，使视觉被完全遮挡时仍有连续控制能力。训练适应头时必须用学生自身 rollout 数据，避免只在完美教师轨迹上训练。")
add_callout(doc, "复现重点：", "0.5 s 本体历史、隐变量而非参数辨识、学生自身 rollout、不同频率异步推理。")

# Paper 2
doc.add_heading("2.2 Learning Robust Perceptive Locomotion for Quadrupedal Robots in the Wild", level=2)
add_source_line(doc, "T. Miki, J. Lee, J. Hwangbo, L. Wellhausen, V. Koltun, M. Hutter", "2022", "2201.08117", "02_Robust_Perceptive_Locomotion_2201.08117.pdf", "Science Robotics 7(62)")
add_label_para(doc, "方法介绍。", "论文先用特权高程与不可观测状态训练教师，再将其蒸馏到学生。学生以 GRU 信念编码器融合本体和带噪足周高程采样；由中间信念状态生成注意力门 α，逐维控制外感知特征进入最终信念状态的比例。训练同时最小化教师-学生动作误差，以及对无噪高程和特权状态的重建误差。感知噪声被拆成采样点横向偏移、高度扰动、足级偏移、回合级漂移和离群点，并按正常、较大偏移、近似失效三种工况以 60%/30%/10% 采样；噪声强度通过课程逐步增加。消融显示 GRU 优于 MLP，带门控优于无门控。")
add_label_para(doc, "我的想法。", "这是与开题报告最接近、最应该作为第一基线的论文。它已经实现了“视觉好时利用视觉、视觉坏时平滑退化”，因此本课题的创新不能只停留在增加 Cross-Attention。真正可区分之处应是：把可校准的不确定性显式输入门控和奖励塑形，并在人形平台上验证其是否产生风险条件化步态。")
add_label_para(doc, "如何应用。", "直接复用其三类噪声工况与课程比例，扩展为高斯噪声、漂移、局部遮挡、全图丢失、时延和飞点。以该论文的 GRU+gate 作为强基线；提出方法则改为“本体 Query + 带方差的高程 token + 连续视觉残差门”。训练损失保留动作蒸馏、地图/特权重建，并新增风险校准损失。")
add_callout(doc, "对比实验必须包含：", "本体盲态、直接拼接、GRU+注意力门、Cross-Attention 无风险量、完整不确定性感知方法。")

# Paper 3
doc.add_heading("2.3 LocoTransformer：Cross-Modal Transformer for Vision-Guided Locomotion", level=2)
add_source_line(doc, "R. Yang, M. Zhang, N. Hansen, H. Xu, X. Wang", "2022", "2107.03996", "03_LocoTransformer_2107.03996.pdf", "ICLR 2022")
add_label_para(doc, "方法介绍。", "LocoTransformer 用 MLP 编码 93 维本体历史，用 ConvNet 把最近 4 帧 64×64 深度图压缩成 4×4 空间 token，再把一个本体 token 与 16 个视觉 token 送入两层、256 维共享 Transformer。自注意力同时完成跨模态关联和视觉空间选择；为避免视觉 token 数量远多于本体 token 而稀释本体信息，输出阶段先按模态分别池化，再拼接进入策略头与价值头。策略使用 PPO 训练，论文还可视化了本体 token 对障碍物和局部地形区域的注意力。")
add_label_para(doc, "我的想法。", "该论文证明 Cross-Modal Transformer 能让本体状态选择与当前运动相关的地形区域，但其注意力矩阵主要是可解释的相关性，不是可靠性估计。开题报告若直接把注意力熵解释成不确定性，需要通过地图误差、遮挡标签和跌倒风险进行校准验证。")
add_label_para(doc, "如何应用。", "将原始深度图 token 替换为局部高程图 token，每个 token 至少包含高度均值、方差/置信度、最近更新时间和有效点比例。采用本体引导的单向 Cross-Attention：Q 来自本体历史，K/V 来自地形 token，并在注意力 logits 中加入 -λ·u_cell，使高不确定区域被连续抑制。保留模态分别池化，避免视觉 token 数量优势造成信息失衡。")

# Paper 4
doc.add_heading("2.4 Learning Vision-Based Bipedal Locomotion for Challenging Terrain", level=2)
add_source_line(doc, "H. Duan, B. Pandit, M. S. Gadde, B. van Marum, J. Dao, C. Kim, A. Fern", "2024", "2309.14594", "04_Vision_Based_Bipedal_Locomotion_2309.14594.pdf", "ICRA 2024")
add_label_para(doc, "方法介绍。", "论文为双足机器人构建两阶段系统。运动控制器接收本体、指令、周期时钟和机器人局部坐标系下 1.5 m×1.0 m、5 cm 分辨率的前向高程图；策略由预训练盲态 LSTM 和视觉调制器组成，后者输出对盲态动作的残差以及步态时钟增量。PPO 训练使用特权 critic、镜像损失、地形/动力学随机化和最高 100 ms 感知延迟。奖励除基础步态外，加入足端加速度和足尖碰撞惩罚。随后用策略 rollout 生成 3 万条深度-状态-真值高程数据，LSTM 重建时序高程图，再用 U-Net 清理边缘；深度图还注入相机位姿、FOV、Gaussian、旋转、边缘、斑点等噪声。")
add_label_para(doc, "我的想法。", "这篇论文对 Atom-01 的价值非常直接：本地坐标高程图减少对全局里程计的依赖，时延随机化和足尖碰撞奖励对双足实机尤为重要。其盲态主干+视觉残差的物理含义也比完全黑箱拼接更清晰。局限是运行时仍有高程预测器与控制器两个模块，但可以保留地图模块，把动作侧整合为一个 actor。")
add_label_para(doc, "如何应用。", "用 elevation_mapping_cupy 输出局部地图，控制策略采用 a_t=a_prop+g_t⊙Δa_vis：a_prop 来自共享本体主干，Δa_vis 来自跨模态分支，g_t 为连续门控；它仍是单一 actor，没有离散策略切换。仿真必须加入 20-150 ms 视觉时延、整图 XYZ 漂移、逐步漂移、足尖碰撞和足端加速度指标。")

# Paper 5
doc.add_heading("2.5 DreamWaQ：Implicit Terrain Imagination from Proprioception", level=2)
add_source_line(doc, "I. M. A. Nahrendra, B. Yu, H. Myung", "2023", "2301.10602", "05_DreamWaQ_2301.10602.pdf", "ICRA 2023")
add_label_para(doc, "方法介绍。", "DreamWaQ 把腿足运动建模为 POMDP，actor 只依赖本体观测历史，critic 可访问外力和高程图。其 CENet 用共享编码器同时完成机体速度估计和 β-VAE 下一观测重建，从历史中得到身体速度与环境上下文 z；策略由此“想象”不可见地形与扰动。论文还提出功率分布奖励，惩罚各电机功率的不均衡；AdaBoot 根据域随机化环境回报的变异系数自动调节估计器输出参与策略训练的概率，避免早期估计噪声破坏学习。")
add_label_para(doc, "我的想法。", "DreamWaQ 对本课题的关键价值不是替代视觉，而是提供一个独立的动力学预测通道。视觉高程预测与本体内部模型对下一状态的预测若长期冲突，往往意味着视觉漂移、地图过期或未知接触，这比单独观察注意力熵更接近真正的跨模态不一致。")
add_label_para(doc, "如何应用。", "增加轻量辅助头预测下一时刻本体状态或机体速度，计算 e_dyn=||o_{t+1}-ô_{t+1}||。将 e_dyn 与地图方差、缺失率组合成风险量。训练早期用 AdaBoot 思路逐步增加辅助估计对策略的影响，避免未收敛风险头反向干扰 PPO。功率分布奖励可直接用于 Atom-01，降低单关节过热风险。")

# Paper 6
doc.add_heading("2.6 Learning Humanoid Locomotion with Perceptive Internal Model", level=2)
add_source_line(doc, "J. Long, J. Ren, M. Shi, Z. Wang, T. Huang, P. Luo, J. Pang", "2025", "2411.14386", "06_Perceptive_Internal_Model_2411.14386.pdf", "ICRA 2025")
add_label_para(doc, "方法介绍。", "PIM 面向人形机器人，以机器人中心、重力对齐坐标系采样 0.8 m×1.2 m 范围内 96 个高程点。内部模型把本体历史与当前高程共同编码，预测下一步状态、机体线速度和环境潜变量；目标编码器提供下一状态表征，通过对比/对齐损失训练内部模型，PPO 更新策略时冻结内部模型，随后再用收集轨迹更新内部模型。系统还使用左右对称正则、逐步开放上肢/腰部关节的动作课程，以及足距、足底平行、接触冲量等人形专用奖励，并在 Unitree H1 和 Fourier GR-1 上验证。")
add_label_para(doc, "我的想法。", "这是十篇中与“人形+高程图+内部模型”最直接的参考。PIM 说明感知不仅可用于落足规划，还能改善状态估计。课题应把 PIM 作为人形基线之一，并把内部模型预测残差纳入不确定性，而不是只在视觉编码层计算一个熵。")
add_label_para(doc, "如何应用。", "Atom-01 可从 96 点采样或规则小网格开始，先复现 PIM 的本体+高程内部模型；然后替换为带方差通道的跨模态注意力。奖励中加入双足横向距离、足底平行度、接触冲量、关节限位和动作二阶平滑。使用左右镜像增强，降低样本量并改善步态对称性。")

# Paper 7
doc.add_heading("2.7 No More Blind Spots：Omnidirectional Vision-Based Bipedal Locomotion", level=2)
add_source_line(doc, "M. S. Gadde, P. Dugar, A. Malik, A. Fern", "2025", "2508.11929", "07_No_More_Blind_Spots_2508.11929.pdf", "Humanoids 2025")
add_label_para(doc, "方法介绍。", "论文以冻结的盲态 LSTM 提供基础动作，视觉策略输出差分动作。教师接收机器人周围特权高程图；学生用四个方向的深度相机，每幅图经共享 ResNet-18 编码，再由两层 LSTM 融合本体、指令、视觉和盲态动作。学生以 DAgger rollout 收集数据，最小化教师-学生动作误差和视觉表征误差。为降低深度渲染成本，论文对同一状态/图像替换多组速度指令，重新查询教师以扩增监督数据，并注入 Gaussian、dropout 和抖动噪声。")
add_label_para(doc, "我的想法。", "四相机全向方案对单前视相机的 Atom-01 未必可直接复制，但差分动作、DAgger 和命令重采样非常有用。命令重采样能在不增加渲染的情况下覆盖不同速度/转向意图，尤其适合 Isaac Lab 中视觉渲染昂贵的学生训练。")
add_label_para(doc, "如何应用。", "若 Atom-01 当前只有前视深度相机，先限制高难地形的倒退/横移指令，避免不可观测动作；若后续加入侧向相机，再扩展全向。训练学生时复用同一帧高程/深度和状态，随机替换 v_x、v_y、ω_z 指令并由教师重新标注，提升数据效率。差分动作结构可与单 actor 的视觉残差分支完全兼容。")

# Paper 8
doc.add_heading("2.8 AME-2：Attention-Based Neural Map Encoding with Uncertainty", level=2)
add_source_line(doc, "C. Zhang, V. Klemm, F. Yang, M. Hutter", "2026", "2601.08485", "08_AME2_2601.08485.pdf", "arXiv preprint / journal submission")
add_label_para(doc, "方法介绍。", "AME-2 的地图为每个栅格输出 (x,y,z,u)，其中 u 是不确定性。轻量 U-Net 从局部深度点云网格联合预测高程与 log-variance，再用里程计投影到全局地图。其概率 Winner-Take-All 融合不会因重复观察同一遮挡区域而虚假降低方差，并拒绝高方差或与先验不一致的更新。控制器先提取局部栅格特征和全局地形特征，再将“本体嵌入+全局特征”形成 Query，对局部特征做多头注意力；学生使用 20 步本体历史的 LSIO 编码，并以 PPO + 动作蒸馏 + 表征对齐联合训练。消融显示全局上下文、学生 RL 损失和表征损失都显著提升未见地形泛化。")
add_label_para(doc, "我的想法。", "AME-2 是本课题“不确定性感知”最强的直接技术依据。它把不确定性放在地图栅格层，而不是事后从注意力矩阵猜测；并证明缺点/伪点增加时，地图方差上升可使策略自然更保守。建议以此作为不确定性主线，再把注意力熵作为辅助特征。")
add_label_para(doc, "如何应用。", "在 elevation_mapping_cupy 旁增加每格方差、有效点比例与更新时间通道；若不能直接修改地图包，可在策略前单独维护这些统计量。Cross-Attention 的 K/V 使用 [height, variance, age, valid_ratio]，并把高方差写入注意力偏置。训练地图头采用异方差回归损失，并用真实高程误差检验 NLL、覆盖率和校准曲线。保留 AME-2 的 PPO+蒸馏+表征损失组合，而不是只做行为克隆。")
add_callout(doc, "最高优先级建议：", "先实现“高程均值+方差+更新时间”的地图输入，再讨论注意力熵；这会显著提高论文创新的可信度和可验证性。")

# Paper 9
doc.add_heading("2.9 A Hybrid Autoencoder for Robust Heightmap Generation", level=2)
add_source_line(doc, "D. Bank, J. Cordes, T. Seel, S. F. G. Ehlers", "2026", "2602.05855", "09_Hybrid_Autoencoder_2602.05855.pdf", "arXiv preprint")
add_label_para(doc, "方法介绍。", "该论文为人形机器人构建机器人中心高程图，实验证明约 7 cm 栅格、0.98 m×0.70 m 范围并前移 0.2 m 在精度与策略可训练性间较合适。深度与 LiDAR 分别经 CNN 自编码器预训练，各得到 256 维潜变量，再与 15 维机器人状态、上一帧 165 维高程图拼接，经两层 256 单元 GRU 输出当前高程图。第一阶段用带 1 cm Gaussian 噪声与 3% 随机遮挡的传感器数据做无监督重建；第二阶段以仿真真值高程监督时序融合。3.2 s 历史比短窗口显著降低误差，但像素 MSE 会把台阶锐边平滑为坡面。")
add_label_para(doc, "我的想法。", "即使 Atom-01 暂无 LiDAR，这篇论文仍能提供高程图尺寸、前向偏置、上一帧地图回馈和 GRU 时间窗的设计依据。其不足恰好提示：若课题以台阶、梅花桩为核心，单纯像素 MSE 不够，应增加边缘/梯度损失或几何一致性损失。")
add_label_para(doc, "如何应用。", "先做深度单分支版本：CNN 编码当前深度或局部点云，拼接 IMU/姿态和上一高程图，以 GRU 输出均值与 log-variance。建议从 1.0 m×0.7 m、7 cm、前移 0.2 m、10-20 Hz 地图开始；训练损失使用 L1/Huber + 高程梯度损失 + 异方差 NLL，以保护台阶边缘。若后续增加 LiDAR，再恢复双分支融合。")

# Paper 10
doc.add_heading("2.10 MARCH：Model-Assisted RL over Sparse Footholds", level=2)
add_source_line(doc, "C. Crismariu, R. K. Cosner", "2026", "2606.10288", "10_MARCH_2606.10288.pdf", "arXiv preprint")
add_label_para(doc, "方法介绍。", "MARCH 先用简化模型生成安全参考：在约束落足区域内用短视距随机 shooting 求解未来 4 步落足与步时，以 Bezier 曲线生成摆脚轨迹，以 HLIP 模型生成质心参考，并约束骨盆、摆脚姿态和上身。教师 PPO 访问这些参考和真值地形，利用 V(η)=1/2·ηᵀPη 构造两项 CLF 启发奖励：奖励跟踪误差小，以及 Lyapunov 候选量随时间下降。学生只接收 32×24 深度图与 5 步本体历史，采用 CNN+Transformer 融合，并用混合密度网络输出双峰动作分布，以处理左右脚或不同落足点的多模态决策。")
add_label_para(doc, "我的想法。", "该方法很适合梅花桩和稀疏落足点，且能为“安全保守行为”提供比手工奖励更强的结构先验。不过它仍是 2026 年预印本，且 CLF 只作为奖励启发，不构成严格闭环安全证明。因此建议将其作为教师或奖励塑形扩展，而不是课题主线的必要依赖。")
add_label_para(doc, "如何应用。", "在梅花桩实验中，用简单落足规划器离线生成安全参考，只供 critic/教师和 CLF 奖励使用；运行时 actor 仍仅使用本体、高程图与不确定性。将风险量 ρ 与 CLF 奖励结合：感知越不确定，越强调参考附近的稳定状态、动作平滑和足端净空。若策略在左右落足选择上出现平均化，可尝试 MDN 或离散落足模式头。")

doc.add_page_break()
doc.add_heading("3. 建议的统一方法：不确定性条件化的单体跨模态策略", level=1)
add_para(doc, "综合十篇论文，建议把研究目标明确为：训练阶段允许特权教师、非对称 critic、地图监督与辅助内部模型；部署阶段只保留一个 actor。actor 同时具有强本体主干、带置信度的地形分支和连续视觉残差门，在任何感知质量下都输出连续动作。")

doc.add_heading("3.1 输入与表征", level=2)
add_bullet(doc, "本体输入：IMU 投影重力、角速度、关节角/速度、上一动作、足端接触和速度指令，堆叠 10-20 个控制步；用 TCN/GRU 得到 z_prop。")
add_bullet(doc, "地形输入：局部高程均值 μ_h、逐格方差 σ²_h、有效点比例、距上次更新的时间 age；由 CNN 生成地形 token。")
add_bullet(doc, "内部模型：用本体历史和地形 token 预测下一本体状态，得到跨模态创新残差 e_dyn。")
add_bullet(doc, "风险标量：ρ_t=sigmoid(b0+b1·mean(σ²_h)+b2·e_dyn+b3·m_drop+b4·H_attn)。其中 H_attn 仅为辅助量，主项是地图方差、时序残差和缺失率。")

doc.add_heading("3.2 跨模态融合与动作输出", level=2)
add_para(doc, "用本体状态作为 Query、地形 token 作为 Key/Value。对第 i 个栅格，将注意力 logit 改为 s_i=(qᵀk_i)/√d-λ_u·u_i-λ_a·age_i，使不可靠或过期区域连续降权。动作采用残差形式：a_t=a_prop+g_t⊙Δa_vis，其中 a_prop 由本体主干产生，Δa_vis 由融合特征产生，g_t=sigmoid(W[z_prop,z_map,ρ_t])。当视觉完整时 g_t 增大；遮挡或地图过期时 g_t 平滑下降到接近 0，而不是切换到另一个策略。")
add_callout(doc, "关键定义：", "这仍是单一 actor。a_prop 和 Δa_vis 是同一网络中的两条连续计算路径，参数可共享，门控连续可微，不存在外部评估器驱动的离散硬切换。")

doc.add_heading("3.3 不确定性条件化的奖励塑形", level=2)
add_para(doc, "建议先定义经监督校准的风险 ρ_t，再在训练时动态调整奖励系数：速度跟踪权重 w_v(ρ)=w_v0·(1-ρ)^γ；足端净空权重 w_c(ρ)=w_c0+k_cρ；姿态稳定和动作平滑权重分别随 ρ 增加。完整奖励可写为：")
add_callout(doc, "建议形式：", "r_t=w_v(ρ)r_vel+w_yaw r_yaw+w_c(ρ)r_clear+w_s(ρ)r_stable+w_a(ρ)r_smooth+r_contact+r_energy+r_survival。", fill="EEF4FA")
add_para(doc, "必须同时把 ρ_t 输入 actor，否则训练期的动态权重未必在推理期产生因果一致的行为。为防止策略通过故意制造高不确定性逃避速度任务，可对风险估计器使用停止梯度，并加入风险校准损失和最低任务完成约束。")

doc.add_heading("3.4 建议的五阶段训练流程", level=2)
for number, text in enumerate([
    "盲态预训练：在平地、坡地和轻度扰动下训练本体主干，确保无视觉时仍可站立、行走和抗推。",
    "特权教师：actor/critic 使用真值高程和动力学参数，训练地形适应与速度跟踪；可加入 PIM/RMA 辅助隐变量。",
    "可部署学生：输入含噪地图与本体历史，使用 PPO + 动作蒸馏 + 表征对齐 + 下一状态预测联合训练。",
    "对抗性感知退化课程：从 nominal 开始，逐步增加 offset、holes、outliers、50-200 ms delay、frame drop 和 blackout；建议初始采用 60/30/10 的正常/偏移/严重失效比例，再依据失败率自适应调整。",
    "Sim-to-Sim 与实机微调前验证：固定策略迁移 MuJoCo；通过后再导出 ONNX/TensorRT，在 Atom-01 上按站立、低速、简单台阶、复杂地形、感知失效逐级测试。",
], 1):
    add_number(doc, text, number)

doc.add_heading("3.5 建议的消融实验与评价指标", level=2)
add_label_para(doc, "基线。", "P0 本体盲态；P1 视觉/本体直接拼接；P2 Miki 式 GRU+门控；P3 Cross-Attention 但无显式不确定性；P4 固定奖励+不确定性输入；P5 动态奖励但不输入风险；P6 完整方法；另设传统硬切换系统作为控制连续性对照。")
add_label_para(doc, "感知退化工况。", "Gaussian 噪声、逐格飞点、整图漂移、局部遮挡、随机丢帧、50/100/150/200 ms 时延、相机完全遮挡、地图冻结，以及地形突变与感知失效同时发生。")
add_label_para(doc, "运动指标。", "通过率/跌倒率、速度跟踪 RMSE、姿态角峰值、足尖碰撞率、足端净空、关节力矩峰值与变化率、动作 jerk、能耗/CoT、恢复时间、控制频率和端到端时延。")
add_label_para(doc, "不确定性指标。", "高程 MAE/RMSE、异方差 NLL、置信区间覆盖率、校准误差、感知失效检测 AUROC、风险与未来 0.5-1.0 s 跌倒/碰撞事件的相关性。只有这些指标成立后，才能宣称风险量真正反映感知不确定性。")

doc.add_page_break()
doc.add_heading("4. 面向 Atom-01 的实施顺序", level=1)
add_callout(doc, "第一优先级（先跑通）：", "复现纯本体 PPO + 局部真值高程图策略；建立 nominal/offset/blackout 三类感知工况；完成 Miki 式门控基线。")
add_callout(doc, "第二优先级（形成创新）：", "在 elevation_mapping_cupy 输出旁增加方差/有效率/更新时间；实现本体 Query 的 Cross-Attention；构建经监督校准的 ρ_t，并把它同时用于门控与奖励条件化。", fill="EEF4FA")
add_callout(doc, "第三优先级（提高论文质量）：", "加入 PIM/DreamWaQ 下一状态预测、RMA 本体隐变量、完整消融和不确定性校准；开展 MuJoCo Sim-to-Sim。")
add_callout(doc, "第四优先级（实机扩展）：", "ONNX/TensorRT 优化、ROS2 时钟与缓存管理、渐进式实机测试；如有余力，再加入 MARCH 的 CLF 教师塑形或多传感器高程融合。")

doc.add_heading("4.1 最小可行网络配置", level=2)
table = doc.add_table(rows=1, cols=3)
for i, h in enumerate(["模块", "建议起始配置", "说明"]):
    set_cell_shading(table.rows[0].cells[i], LIGHT)
    p = table.rows[0].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(h), east="微软雅黑", size=9.5, bold=True, color=DARK)
repeat_header(table.rows[0])
configs = [
    ("本体编码", "20步历史；TCN/GRU；128维", "兼顾 RMA/PIM 式动态表征"),
    ("高程图", "约1.0m×0.7m；7cm；前移0.2m", "从约15×11栅格起步"),
    ("地图通道", "高度、方差、age、valid ratio", "至少4通道"),
    ("融合", "2层 Cross-Attention；4头；128维", "本体为 Query"),
    ("动作", "关节目标位置+共享本体主干的视觉残差", "外部 PD 跟踪"),
    ("辅助头", "下一本体状态、风险、地图重建", "训练期使用，可部署裁剪"),
    ("控制频率", "actor 50-100Hz；地图10-20Hz", "缓存最近可靠地图与 age"),
]
for row in configs:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        p = cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(p.add_run(value), size=9.3)
set_table_geometry(table, [1700, 3300, 4300], indent=120)
set_table_borders(table)

doc.add_heading("4.2 近期最值得先做的三个实验", level=2)
add_number(doc, "验证注意力熵：固定机器人状态，对高程图逐步加入漂移、遮挡和随机噪声，比较 H_attn 与真实地图误差、方差和未来碰撞率的相关性。若相关性不稳定，不再把它作为唯一不确定性。", 1)
add_number(doc, "验证风险条件化行为：在相同台阶上只改变感知质量，检查 ρ 上升是否稳定导致期望速度下降、足端净空上升、动作 jerk 下降，并确认不存在故意停滞或风险估计作弊。", 2)
add_number(doc, "验证连续软退化：让视觉质量在临界区间反复变化，比较完整方法与硬切换基线的动作差、关节力矩差和姿态波动，直接证明单网络连续门控消除了切换抖动。", 3)

doc.add_heading("5. 本地论文文件与公开来源", level=1)
add_para(doc, "以下 10 篇均已下载为作者公开的 arXiv PDF，并与本节顺序一致。开题报告原文件保留不变。", first_indent=False)

papers = [
    ("01_RMA_2107.04034.pdf", "RMA: Rapid Motor Adaptation for Legged Robots", "2107.04034"),
    ("02_Robust_Perceptive_Locomotion_2201.08117.pdf", "Learning Robust Perceptive Locomotion for Quadrupedal Robots in the Wild", "2201.08117"),
    ("03_LocoTransformer_2107.03996.pdf", "Learning Vision-Guided Quadrupedal Locomotion End-to-End with Cross-Modal Transformers", "2107.03996"),
    ("04_Vision_Based_Bipedal_Locomotion_2309.14594.pdf", "Learning Vision-Based Bipedal Locomotion for Challenging Terrain", "2309.14594"),
    ("05_DreamWaQ_2301.10602.pdf", "DreamWaQ: Learning Robust Quadrupedal Locomotion with Implicit Terrain Imagination", "2301.10602"),
    ("06_Perceptive_Internal_Model_2411.14386.pdf", "Learning Humanoid Locomotion with Perceptive Internal Model", "2411.14386"),
    ("07_No_More_Blind_Spots_2508.11929.pdf", "No More Blind Spots: Learning Vision-Based Omnidirectional Bipedal Locomotion", "2508.11929"),
    ("08_AME2_2601.08485.pdf", "AME-2: Agile and Generalized Legged Locomotion via Attention-Based Neural Map Encoding", "2601.08485"),
    ("09_Hybrid_Autoencoder_2602.05855.pdf", "A Hybrid Autoencoder for Robust Heightmap Generation from Fused Lidar and Depth Data", "2602.05855"),
    ("10_MARCH_2606.10288.pdf", "MARCH: Model-Assisted Reinforcement Learning for the Perceptive Control of Humanoids", "2606.10288"),
]
for i, (fname, title, aid) in enumerate(papers, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    set_run_font(p.add_run(f"{i}. "), size=9.8)
    set_run_font(p.add_run(f"{title}。"), bold=True, size=9.8)
    set_run_font(p.add_run(f" 本地：{fname}；公开来源："), size=9.8)
    add_hyperlink(p, f"https://arxiv.org/abs/{aid}", f"https://arxiv.org/abs/{aid}")

doc.add_heading("6. 最终建议", level=1)
add_para(doc, "本课题最有潜力的论文主张应是：在单一人形机器人控制策略中，把经过校准的逐格感知不确定性和跨模态动力学不一致显式纳入视觉-本体连续融合，并通过风险条件化训练使策略在感知退化时主动形成可测量的保守步态。Cross-Attention 是实现手段，不确定性校准、推理期风险输入、连续动作退化和系统化消融才是论证创新与有效性的关键。")
add_callout(doc, "推荐的论文贡献表述：", "提出一种不确定性条件化的本体引导跨模态策略；提出地图方差与内部模型创新融合的风险估计；提出风险条件化奖励与对抗性感知退化课程；在 Atom-01 上证明其相对直接拼接、门控融合和硬切换方法具有更低跌倒率、更小力矩突变和更平滑的视觉失效退化。", fill="EEF4FA")

# Document metadata
props = doc.core_properties
props.title = "基于不确定性感知的人形机器人鲁棒自适应运动控制：首批10篇论文方法综述"
props.subject = "论文方法介绍、研究判断与 Atom-01 应用方案"
props.author = "Codex（依据林金诚开题报告整理）"
props.keywords = "humanoid locomotion, uncertainty, cross-modal attention, reinforcement learning, Atom-01"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
